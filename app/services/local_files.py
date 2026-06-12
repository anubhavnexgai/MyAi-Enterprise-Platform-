"""Local file access for the assistant — find / read / open the user's own files
by natural language ("open the latest PRD for MyAi").

Scoped to the user's common folders (Documents, Desktop, Downloads, OneDrive,
MyAiProjects, Pictures) and the home directory — never system folders. The agent
does the natural-language reasoning: it calls ``find_files`` with keywords, gets
candidates with modified-times, and picks the right one (e.g. the most recent).
"""
from __future__ import annotations

import os
import re
import time
from pathlib import Path
from typing import Dict, List

# Filler words to drop from a natural-language file query.
_STOP = {
    "the", "a", "an", "my", "latest", "newest", "recent", "last", "open", "file",
    "files", "for", "of", "find", "show", "please", "document", "doc", "to", "in",
    "on", "and", "me", "that", "this", "named", "called", "named",
}
_SKIP_DIRS = {
    "node_modules", ".git", "__pycache__", ".venv", "venv", "env", "appdata",
    "$recycle.bin", "windows", "program files", "program files (x86)", "programdata",
    ".cache", "site-packages", ".next", "dist", "build", ".idea", ".vscode",
}
_TEXT_EXT = {
    ".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".csv", ".html",
    ".css", ".yaml", ".yml", ".log", ".ini", ".cfg", ".xml", ".sh", ".bat", ".sql",
    ".env", ".toml", ".rst",
}


def _roots() -> List[Path]:
    home = Path.home()
    cands = [home / "Documents", home / "Desktop", home / "Downloads",
             home / "OneDrive", home / "MyAiProjects", home / "Pictures"]
    seen, out = set(), []
    for p in cands:
        rp = p.resolve()
        if p.exists() and str(rp) not in seen:
            seen.add(str(rp))
            out.append(p)
    return out


def _allowed(p: Path) -> bool:
    try:
        rp = str(p.resolve())
    except Exception:  # noqa: BLE001
        return False
    home = str(Path.home().resolve())
    return rp.startswith(home)


def find_files(query: str, limit: int = 12) -> List[Dict]:
    """Find files in the user's folders matching a natural-language query. Returns
    the best matches sorted by relevance then most-recently-modified, so 'latest X'
    naturally lands first."""
    terms = [t for t in re.split(r"\W+", (query or "").lower()) if t and t not in _STOP]
    results: List[dict] = []
    for root in _roots():
        for dirpath, dirnames, filenames in os.walk(root):
            try:
                depth = len(Path(dirpath).relative_to(root).parts)
            except Exception:  # noqa: BLE001
                depth = 0
            if depth > 5:
                dirnames[:] = []
                continue
            dirnames[:] = [d for d in dirnames
                           if d.lower() not in _SKIP_DIRS and not d.startswith(".")]
            for fn in filenames:
                low = fn.lower()
                score = sum(1 for t in terms if t in low)
                if terms and score == 0:
                    continue
                p = Path(dirpath) / fn
                try:
                    st = p.stat()
                except Exception:  # noqa: BLE001
                    continue
                results.append({"path": str(p), "name": fn, "dir": dirpath,
                                "score": score, "mtime": st.st_mtime, "size": st.st_size})
            if len(results) > 4000:
                break
    results.sort(key=lambda r: (r["score"], r["mtime"]), reverse=True)
    out = []
    for r in results[:limit]:
        out.append({
            "path": r["path"], "name": r["name"], "folder": r["dir"],
            "modified": time.strftime("%Y-%m-%d %H:%M", time.localtime(r["mtime"])),
            "size_kb": round(r["size"] / 1024, 1),
        })
    return out


def _extract_docx(p: Path) -> str:
    """Pull text out of a .docx (paragraphs + tables)."""
    import docx  # python-docx
    d = docx.Document(str(p))
    parts = [para.text for para in d.paragraphs if para.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(p: Path, max_pages: int = 40) -> str:
    """Pull text out of a .pdf (first ``max_pages`` pages)."""
    from pypdf import PdfReader
    reader = PdfReader(str(p))
    out = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            out.append(f"…(stopped at {max_pages} pages)")
            break
        try:
            out.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    return "\n".join(out)


def read_file(path: str, max_chars: int = 12000) -> str:
    """Read a file's text contents (within the user's folders). Handles plain
    text/markdown/code, plus .docx (Word) and .pdf via extraction."""
    p = Path(path)
    if not _allowed(p):
        return "BLOCKED: that path is outside your home folder."
    if not p.is_file():
        return f"No file found at {path}."
    ext = p.suffix.lower()
    try:
        if ext == ".docx":
            txt = _extract_docx(p)
        elif ext == ".pdf":
            txt = _extract_pdf(p)
        elif ext in _TEXT_EXT:
            txt = p.read_text(encoding="utf-8", errors="replace")
        else:
            return (f"'{p.name}' is a {ext or 'binary'} file I can't read as text — "
                    "I can open it in its app with open_file.")
    except Exception as exc:  # noqa: BLE001
        return f"Could not read {p.name}: {exc}"
    if not txt.strip():
        return f"'{p.name}' opened but has no extractable text (it may be scanned/image-only)."
    return txt[:max_chars] + ("\n…(truncated)" if len(txt) > max_chars else "")


def open_file(path: str) -> str:
    """Open a file with its default app on the user's machine."""
    p = Path(path)
    if not _allowed(p):
        return "BLOCKED: that path is outside your home folder."
    if not p.exists():
        return f"No file found at {path}."
    try:
        os.startfile(str(p))  # Windows: launch with the default associated app
    except AttributeError:
        import subprocess
        subprocess.Popen(["xdg-open", str(p)])  # non-Windows fallback
    except Exception as exc:  # noqa: BLE001
        return f"Could not open {path}: {exc}"
    return f"Opened '{p.name}' in its default app."
